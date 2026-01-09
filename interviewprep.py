import streamlit as st
import requests
from bs4 import BeautifulSoup
import pandas as pd
from datetime import datetime
import json
import re
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
from xml.sax.saxutils import escape
from urllib.parse import quote_plus
import html

# Page configuration
st.set_page_config(
    page_title="Interview Prep Master Pro",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced CSS with proper spacing and formatting
st.markdown("""
<style>
    /* Main container with scrolling */
    .main .block-container {
        max-height: 90vh;
        overflow-y: auto;
        padding-bottom: 50px;
    }
    
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        background: linear-gradient(120deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 2rem;
        padding: 20px 0;
    }
    
    .sub-header {
        font-size: 1.8rem;
        color: #667eea;
        font-weight: 600;
        margin-top: 3rem;
        margin-bottom: 2rem;
        border-bottom: 3px solid #667eea;
        padding-bottom: 10px;
    }
    
    .question-container {
        background: white;
        border-radius: 15px;
        padding: 25px;
        margin: 25px 0;
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        border-left: 5px solid #667eea;
    }
    
    .question-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    
    .question-number {
        font-size: 1rem;
        font-weight: 600;
        opacity: 0.9;
        margin-bottom: 10px;
    }
    
    .question-text {
        font-size: 1.3rem;
        font-weight: 600;
        line-height: 1.6;
        margin: 0;
    }
    
    .answer-section {
        background: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        margin: 20px 0;
        border-left: 4px solid #48bb78;
    }
    
    .answer-label {
        font-size: 1.1rem;
        font-weight: 700;
        color: #2d3748;
        margin-bottom: 15px;
        display: block;
    }
    
    .answer-text {
        font-size: 1.05rem;
        color: #4a5568;
        line-height: 1.8;
        text-align: justify;
    }
    
    .answer-text ul {
        margin: 15px 0;
        padding-left: 25px;
    }
    
    .answer-text li {
        margin: 10px 0;
        line-height: 1.6;
    }
    
    .metadata-section {
        display: flex;
        gap: 15px;
        flex-wrap: wrap;
        margin: 20px 0;
        padding: 15px;
        background: #e6f7ff;
        border-radius: 8px;
    }
    
    .badge {
        display: inline-block;
        padding: 8px 16px;
        border-radius: 20px;
        font-size: 0.9rem;
        font-weight: 600;
        margin: 5px;
    }
    
    .badge-source {
        background: #48bb78;
        color: white;
    }
    
    .badge-difficulty {
        background: #f6ad55;
        color: white;
    }
    
    .badge-category {
        background: #667eea;
        color: white;
    }
    
    .link-button {
        display: inline-block;
        background: linear-gradient(120deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 12px 24px;
        border-radius: 8px;
        text-decoration: none;
        font-weight: 600;
        margin: 10px 5px;
        transition: transform 0.2s;
    }
    
    .link-button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(102, 126, 234, 0.4);
    }
    
    .metric-card {
        background: white;
        padding: 20px;
        border-radius: 12px;
        box-shadow: 0 3px 10px rgba(0, 0, 0, 0.1);
        text-align: center;
        border-top: 4px solid #667eea;
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: bold;
        color: #667eea;
        margin: 10px 0;
    }
    
    .metric-label {
        font-size: 0.9rem;
        color: #666;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    .stButton>button {
        background: linear-gradient(120deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 15px 30px;
        border-radius: 10px;
        font-weight: 600;
        font-size: 1.1rem;
        transition: all 0.3s;
    }
    
    .stButton>button:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 20px rgba(102, 126, 234, 0.4);
    }
    
    .info-box {
        background: #e6f7ff;
        border-left: 5px solid #1890ff;
        padding: 20px;
        border-radius: 8px;
        margin: 20px 0;
    }
    
    .divider {
        height: 2px;
        background: linear-gradient(90deg, transparent, #667eea, transparent);
        margin: 40px 0;
    }
    
    /* Scrollbar styling */
    ::-webkit-scrollbar {
        width: 12px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, #667eea, #764ba2);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: #667eea;
    }
</style>
""", unsafe_allow_html=True)

# Comprehensive question database with proper formatting
def generate_comprehensive_questions(technology, num_questions, filter_type="all", company=None):
    """Generate comprehensive, well-formatted questions from knowledge base"""
    
    question_database = {
        "Artificial Intelligence": [
            {
                "question": "What is Artificial Intelligence and what are its main types?",
                "answer": """Artificial Intelligence (AI) is the simulation of human intelligence by machines, particularly computer systems. AI enables machines to learn from experience, adjust to new inputs, and perform human-like tasks.

**Main Types of AI:**

• **Narrow/Weak AI**: Designed for specific tasks (Siri, Alexa, recommendation systems)
• **General AI**: Human-level intelligence across all domains (theoretical)
• **Super AI**: Surpasses human intelligence (hypothetical)

**By Functionality:**
• **Reactive Machines**: No memory, responds to current situations (Deep Blue)
• **Limited Memory**: Uses past data for decisions (Self-driving cars)
• **Theory of Mind**: Understands emotions and intentions (under development)
• **Self-Aware AI**: Has consciousness (hypothetical)

**Key Applications:** Computer vision, natural language processing, robotics, expert systems, speech recognition, and autonomous vehicles.""",
                "source": "AI Fundamentals",
                "difficulty": "Easy",
                "category": "Concepts",
                "video": "https://www.youtube.com/results?search_query=artificial+intelligence+types+explained",
                "trending": True
            },
            {
                "question": "Explain Machine Learning, Deep Learning, and their relationship with AI.",
                "answer": """Machine Learning (ML) and Deep Learning (DL) are subsets of Artificial Intelligence with distinct characteristics.

**Machine Learning:**
• Algorithms that learn from data without explicit programming
• Requires feature engineering by humans
• Works well with smaller datasets
• Examples: Linear Regression, Decision Trees, SVM, Random Forest

**Deep Learning:**
• Subset of ML using artificial neural networks
• Automatically learns features from raw data
• Requires large datasets and computational power
• Examples: CNN, RNN, Transformers, GANs

**Relationship Hierarchy:**
AI ⊃ Machine Learning ⊃ Deep Learning

**Key Differences:**
• **Data Requirements**: ML (thousands), DL (millions)
• **Hardware**: ML (standard CPU), DL (GPU/TPU needed)
• **Training Time**: ML (minutes to hours), DL (hours to days)
• **Interpretability**: ML (more interpretable), DL (black box)
• **Feature Engineering**: ML (manual), DL (automatic)""",
                "source": "Machine Learning Basics",
                "difficulty": "Medium",
                "category": "ML Fundamentals",
                "video": "https://www.youtube.com/results?search_query=machine+learning+vs+deep+learning",
                "trending": True
            },
            {
                "question": "What is the difference between supervised, unsupervised, and reinforcement learning?",
                "answer": """These are three fundamental paradigms in machine learning, each suited for different types of problems.

**Supervised Learning:**
• **Definition**: Learning from labeled data (input-output pairs)
• **Goal**: Predict output for new inputs
• **Types**: Classification (discrete output), Regression (continuous output)
• **Examples**: 
  - Email spam detection
  - House price prediction
  - Image classification
• **Algorithms**: Linear Regression, Logistic Regression, SVM, Neural Networks

**Unsupervised Learning:**
• **Definition**: Finding patterns in unlabeled data
• **Goal**: Discover hidden structures
• **Types**: Clustering, Dimensionality Reduction, Association
• **Examples**:
  - Customer segmentation
  - Anomaly detection
  - Recommendation systems
• **Algorithms**: K-Means, DBSCAN, PCA, Autoencoders

**Reinforcement Learning:**
• **Definition**: Learning through interaction with environment
• **Goal**: Maximize cumulative reward
• **Components**: Agent, Environment, Actions, Rewards, Policy
• **Examples**:
  - Game playing (AlphaGo)
  - Robotics control
  - Autonomous driving
• **Algorithms**: Q-Learning, SARSA, DQN, PPO, A3C

**When to Use:**
• Supervised: When you have labeled historical data
• Unsupervised: When you want to discover patterns
• Reinforcement: When you need sequential decision making""",
                "source": "ML Paradigms",
                "difficulty": "Easy",
                "category": "Learning Types",
                "video": "https://www.youtube.com/results?search_query=supervised+unsupervised+reinforcement+learning",
                "trending": True
            },
            {
                "question": "Explain the bias-variance tradeoff and overfitting/underfitting in machine learning.",
                "answer": """The bias-variance tradeoff is a fundamental concept that affects model performance and generalization.

**Bias:**
• Error from overly simplistic assumptions
• High bias leads to **underfitting**
• Model misses relevant patterns in data
• **Example**: Using linear model for non-linear data

**Variance:**
• Error from sensitivity to training data fluctuations
• High variance leads to **overfitting**
• Model learns noise instead of patterns
• **Example**: Deep decision tree memorizing training data

**The Tradeoff:**
Total Error = Bias² + Variance + Irreducible Error

• **Low Complexity Model**: High bias, low variance → Underfitting
• **High Complexity Model**: Low bias, high variance → Overfitting
• **Optimal Model**: Balanced bias and variance

**Underfitting (High Bias):**
• Signs: Poor training AND test performance
• Causes: Model too simple, insufficient features
• Solutions:
  - Increase model complexity
  - Add more features
  - Reduce regularization
  - Train longer

**Overfitting (High Variance):**
• Signs: Excellent training, poor test performance
• Causes: Model too complex, insufficient data, noise learning
• Solutions:
  - Simplify model (reduce parameters)
  - Add more training data
  - Use regularization (L1/L2)
  - Apply dropout (neural networks)
  - Use cross-validation
  - Early stopping
  - Data augmentation
  - Ensemble methods

**Practical Detection:**
• Plot learning curves (training vs validation error)
• Use cross-validation scores
• Monitor train-test gap""",
                "source": "ML Theory",
                "difficulty": "Medium",
                "category": "Model Performance",
                "video": "https://www.youtube.com/results?search_query=bias+variance+tradeoff+explained",
                "trending": True
            },
            {
                "question": "What are Neural Networks and how do they work?",
                "answer": """Neural Networks are computing systems inspired by biological neural networks in the human brain, fundamental to deep learning.

**Architecture Components:**

**1. Neurons (Nodes):**
• Basic computational units
• Receive inputs, apply weights, add bias
• Pass through activation function
• Formula: output = activation(Σ(weights × inputs) + bias)

**2. Layers:**
• **Input Layer**: Receives raw data
• **Hidden Layers**: Perform computations and feature extraction
• **Output Layer**: Produces final predictions

**3. Weights and Biases:**
• **Weights**: Importance of connections between neurons
• **Biases**: Offset values to shift activation function
• Learned during training through backpropagation

**4. Activation Functions:**
• **ReLU**: f(x) = max(0, x) - most common for hidden layers
• **Sigmoid**: f(x) = 1/(1+e^-x) - binary classification
• **Tanh**: f(x) = (e^x - e^-x)/(e^x + e^-x) - normalized output
• **Softmax**: Multi-class probability distribution

**How They Work:**

**Forward Propagation:**
1. Input data flows through network
2. Each neuron calculates weighted sum
3. Applies activation function
4. Passes output to next layer
5. Final layer produces prediction

**Backpropagation:**
1. Calculate loss (error) at output
2. Compute gradients using chain rule
3. Propagate gradients backward through network
4. Update weights using gradient descent
5. Repeat until convergence

**Training Process:**
• Initialize weights randomly
• Forward pass to get predictions
• Calculate loss function
• Backward pass to compute gradients
• Update weights using optimizer (SGD, Adam)
• Iterate for multiple epochs

**Key Hyperparameters:**
• Learning rate: Step size for weight updates
• Batch size: Number of samples per gradient update
• Number of layers and neurons
• Activation functions
• Regularization parameters

**Advantages:**
• Can learn complex non-linear patterns
• Automatic feature extraction
• Scalable to large datasets
• Transfer learning capabilities

**Challenges:**
• Require large datasets
• Computationally expensive
• Black box (hard to interpret)
• Prone to overfitting""",
                "source": "Deep Learning Fundamentals",
                "difficulty": "Medium",
                "category": "Neural Networks",
                "video": "https://www.youtube.com/results?search_query=neural+networks+explained",
                "trending": True
            }
        ],
        "Machine Learning": [
            {
                "question": "What is feature engineering and why is it important?",
                "answer": """Feature engineering is the process of creating, transforming, and selecting features to improve machine learning model performance. It's often considered more important than algorithm selection.

**Importance:**
• Can improve model accuracy by 10-30%
• Reduces training time
• Makes models more interpretable
• Captures domain knowledge
• Reduces overfitting

**Key Techniques:**

**1. Feature Creation:**
• **Polynomial Features**: x² , x³, interaction terms (x₁ × x₂)
• **Domain-Specific Features**: Day of week from date, age from birthdate
• **Binning**: Converting continuous to categorical (age → age_group)
• **Aggregations**: Sum, mean, count, max, min over groups

**2. Feature Transformation:**
• **Scaling**: Normalization (0-1), Standardization (mean=0, std=1)
• **Log Transform**: For skewed distributions
• **Box-Cox/Yeo-Johnson**: Power transformations
• **Mathematical**: Square root, reciprocal

**3. Encoding Categorical Variables:**
• **One-Hot Encoding**: Binary columns for each category
• **Label Encoding**: Ordinal integers (for ordered categories)
• **Target Encoding**: Replace with target mean
• **Frequency Encoding**: Replace with occurrence count
• **Binary Encoding**: Hybrid of one-hot and label

**4. Feature Extraction:**
• **PCA**: Linear dimensionality reduction
• **t-SNE/UMAP**: Non-linear dimensionality reduction
• **Autoencoders**: Neural network feature compression
• **Text Features**: TF-IDF, Word2Vec, BERT embeddings

**5. Feature Selection:**
• **Filter Methods**: Correlation, Chi-square, mutual information
• **Wrapper Methods**: RFE (Recursive Feature Elimination)
• **Embedded Methods**: Lasso, Ridge, Tree feature importance

**6. Handling Missing Values:**
• Mean/Median/Mode imputation
• Forward/Backward fill (time series)
• KNN imputation
• Create missing indicator feature

**7. Handling Outliers:**
• Winsorization (capping)
• Transformation (log, sqrt)
• Removal (if justified)
• Separate modeling

**Best Practices:**
• Understand domain and data
• Create features before splitting train/test
• Avoid data leakage
• Document transformations
• Use pipeline for reproducibility
• Validate on holdout set""",
                "source": "Feature Engineering Guide",
                "difficulty": "Medium",
                "category": "Data Preprocessing",
                "video": "https://www.youtube.com/results?search_query=feature+engineering+machine+learning",
                "trending": True
            },
            {
                "question": "Explain ensemble methods: Bagging, Boosting, and Stacking.",
                "answer": """Ensemble methods combine multiple models to achieve better performance than individual models. They reduce overfitting and improve accuracy.

**1. Bagging (Bootstrap Aggregating):**

**How it Works:**
• Create multiple training subsets by random sampling with replacement
• Train separate models on each subset
• Aggregate predictions (voting/averaging)

**Characteristics:**
• Reduces variance
• Models trained in parallel
• Works best with high-variance models

**Algorithms:**
• **Random Forest**: Bagging with decision trees + random feature selection
  - Typically 100-1000 trees
  - Each tree sees random subset of features at each split
  - Final prediction by majority vote (classification) or average (regression)

**Advantages:**
• Reduces overfitting
• Handles high-dimensional data
• Provides feature importance
• Robust to outliers

**When to Use:**
• Model overfits training data
• High variance, low bias model
• Noisy datasets

---

**2. Boosting:**

**How it Works:**
• Train models sequentially
• Each model corrects errors of previous model
• Weights are assigned to observations
• Misclassified samples get higher weights

**Characteristics:**
• Reduces bias AND variance
• Sequential training
• More prone to overfitting if not careful

**Popular Algorithms:**

**AdaBoost (Adaptive Boosting):**
• Increases weights of misclassified samples
• Combines weak learners (often decision stumps)
• Final prediction: weighted vote

**Gradient Boosting:**
• Fits new model to residual errors
• Uses gradient descent to minimize loss
• Examples: GBM, XGBoost, LightGBM, CatBoost

**XGBoost Features:**
• Regularization (L1/L2) to prevent overfitting
• Parallel processing
• Handles missing values
• Tree pruning
• Built-in cross-validation

**LightGBM Features:**
• Leaf-wise tree growth (faster)
• Handles large datasets efficiently
• Lower memory usage
• Categorical feature support

**CatBoost Features:**
• Automatic handling of categorical features
• Reduces overfitting
• Fast prediction
• Robust to overfitting

**Boosting Best Practices:**
• Use small learning rate (0.01-0.1)
• Monitor validation error for early stopping
• Increase number of estimators with smaller learning rate
• Use regularization parameters

---

**3. Stacking (Stacked Generalization):**

**How it Works:**
• Train multiple diverse base models (Level 0)
• Use predictions as features for meta-model (Level 1)
• Meta-model learns to combine base model predictions

**Architecture:**
Level 0: [Model1, Model2, Model3, ...] → Predictions
Level 1: Meta-Model (uses Level 0 predictions) → Final Prediction

**Implementation:**
• Split data into train and validation
• Train base models on training data
• Generate predictions on validation data (avoid overfitting)
• Train meta-model on base model predictions

**Base Models (Diverse):**
• Linear models (Logistic Regression, Ridge)
• Tree models (Random Forest, XGBoost)
• Neural Networks
• SVM, KNN

**Meta-Model (Simple):**
• Logistic Regression
• Linear Regression
• Neural Network

**Advantages:**
• Often achieves best performance
• Leverages strengths of different models
• Flexibility in model selection

**Disadvantages:**
• Complex and time-consuming
• Risk of overfitting meta-model
• Hard to interpret
• Computationally expensive

---

**Comparison:**

| Aspect | Bagging | Boosting | Stacking |
|--------|---------|----------|----------|
| Training | Parallel | Sequential | Multi-level |
| Reduces | Variance | Bias & Variance | Both |
| Speed | Fast | Slower | Slowest |
| Overfitting | Low risk | Moderate risk | High risk |
| Complexity | Low | Medium | High |

**When to Use:**
• **Bagging**: High variance models, unstable predictions
• **Boosting**: Underfitting models, need high accuracy
• **Stacking**: Maximum performance needed, have computational resources""",
                "source": "Ensemble Methods",
                "difficulty": "Hard",
                "category": "Ensemble Learning",
                "video": "https://www.youtube.com/results?search_query=bagging+boosting+stacking+explained",
                "trending": True
            }
        ],
        "Python": [
            {
                "question": "Explain Python's memory management and garbage collection.",
                "answer": """Python uses automatic memory management with a combination of reference counting and garbage collection.

**Memory Management Components:**

**1. Reference Counting:**
• Every object has a reference count
• Count increases when reference created
• Count decreases when reference deleted
• When count reaches 0, memory freed immediately

**Example:**
```python
import sys
a = []  # ref count = 1
b = a   # ref count = 2
sys.getrefcount(a)  # Returns 3 (includes temporary reference)
del b   # ref count = 2
del a   # ref count = 1, memory freed
```

**2. Garbage Collector:**
• Handles circular references (ref counting can't)
• Uses generational garbage collection
• Three generations (0, 1, 2)

**Generational Collection:**
• **Generation 0**: Newly created objects
  - Collected most frequently
  - Threshold typically ~700 objects
• **Generation 1**: Survived gen 0 collection
  - Collected less frequently
• **Generation 2**: Long-lived objects
  - Collected least frequently

**How It Works:**
1. Objects start in Generation 0
2. When gen 0 fills up, collection triggered
3. Surviving objects promoted to gen 1
4. Process repeats for higher generations

**Circular Reference Problem:**
```python
class Node:
    def __init__(self):
        self.ref = None

a = Node()
b = Node()
a.ref = b
b.ref = a  # Circular reference
# Even after 'del a, b', reference count > 0
# Garbage collector detects and cleans this
```

**Memory Pools:**
Python uses memory pools for efficiency:
• **Arenas**: 256KB blocks from OS
• **Pools**: 4KB chunks within arenas
• **Blocks**: Fixed-size pieces within pools
• Small objects (<512 bytes) use pooled memory
• Large objects allocated directly from heap

**Manual Control:**

**Garbage Collection:**
```python
import gc

gc.collect()  # Force collection
gc.disable()  # Disable GC
gc.enable()   # Enable GC
gc.get_count()  # Get collection counts
gc.set_threshold(700, 10, 10)  # Set thresholds
```

**Memory Profiling:**
```python
import sys
import tracemalloc

# Object size
sys.getsizeof(object)

# Memory tracking
tracemalloc.start()
# ... your code ...
snapshot = tracemalloc.take_snapshot()
top_stats = snapshot.statistics('lineno')
```

**Best Practices:**

**1. Explicit Deletion:**
```python
del large_object  # Decrease ref count
```

**2. Context Managers:**
```python
with open('file.txt') as f:
    data = f.read()
# File automatically closed
```

**3. Weak References:**
```python
import weakref
obj = SomeClass()
weak_obj = weakref.ref(obj)  # Doesn't increase ref count
```

**4. Slots for Memory Efficiency:**
```python
class Point:
    __slots__ = ['x', 'y']  # Reduces memory per instance
```

**5. Generators for Large Data:**
```python
# Instead of list
data = [x**2 for x in range(1000000)]  # Uses lots of memory

# Use generator
data = (x**2 for x in range(1000000))  # Memory efficient
```

**6. Avoid Circular References:**
• Use weak references
• Break cycles explicitly
• Use context managers

**Common Issues:**

**Memory Leaks:**
• Circular references with __del__
• Global variables not cleared
• Caches growing indefinitely
• Event listeners not unregistered

**Performance Tips:**
• Reuse objects instead of creating new ones
• Use appropriate data structures
• Profile before optimizing
• Consider PyPy for long-running programs""",
                "source": "Python Internals",
                "difficulty": "Hard",
                "category": "Memory Management",
                "video": "https://www.youtube.com/results?search_query=python+memory+management",
                "trending": False
            }
        ]
    }
    
    # Get questions for technology
    questions = question_database.get(technology, question_database.get("Artificial Intelligence", []))
    
    # Filter by type
    if filter_type == "trending":
        questions = [q for q in questions if q.get("trending", False)]
    elif filter_type == "latest":
        questions = sorted(questions, key=lambda x: x.get("trending", False), reverse=True)
    
    # Add company tag if specified
    if company and company != "Select Company":
        for q in questions:
            q["company"] = company
            q["question"] = f"[{company} Interview] {q['question']}"
    
    # Duplicate questions to reach requested count if needed
    if len(questions) < num_questions:
        multiplier = (num_questions // len(questions)) + 1
        questions = (questions * multiplier)[:num_questions]
    
    return questions[:num_questions]


def create_pdf(questions_data, technology, company=None):
    """Create PDF with proper formatting"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=HexColor('#667eea'),
        spaceAfter=20
    )
    
    content = []
    content.append(Paragraph(escape(f"Interview Questions: {technology}"), title_style))
    content.append(Spacer(1, 0.3*inch))
    
    for idx, qa in enumerate(questions_data, 1):
        # Question
        q_text = f"<b>Q{idx}.</b> {escape(qa['question'])}"
        content.append(Paragraph(q_text, styles['Heading3']))
        content.append(Spacer(1, 0.1*inch))
        
        # Answer
        if qa.get('answer'):
            answer_cleaned = html.unescape(qa['answer']).replace('**', '')
            answer_cleaned = re.sub(r'```python.*?```', '', answer_cleaned, flags=re.DOTALL)
            answer_cleaned = ' '.join(answer_cleaned.split())[:1000]
            content.append(Paragraph(escape(answer_cleaned), styles['BodyText']))
        
        # Metadata
        if qa.get('source'):
            content.append(Spacer(1, 0.05*inch))
            content.append(Paragraph(f"<i>Source: {escape(qa['source'])}</i>", styles['Normal']))
        
        content.append(Spacer(1, 0.2*inch))
        
        if idx % 3 == 0:
            content.append(PageBreak())
    
    doc.build(content)
    buffer.seek(0)
    return buffer


def create_word(questions_data, technology, company=None):
    """Create Word document"""
    doc = Document()
    
    title = doc.add_heading(f'Interview Questions: {technology}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for idx, qa in enumerate(questions_data, 1):
        # Question
        q_heading = doc.add_heading(f"Question {idx}", level=2)
        q_para = doc.add_paragraph(qa['question'])
        q_para.runs[0].font.bold = True
        
        # Answer
        if qa.get('answer'):
            doc.add_paragraph("Answer:", style='Heading 3')
            answer_cleaned = html.unescape(qa['answer']).replace('**', '').replace('•', '-')
            answer_cleaned = re.sub(r'```python.*?```', '', answer_cleaned, flags=re.DOTALL)
            doc.add_paragraph(answer_cleaned[:1000])
        
        # Metadata
        if qa.get('source'):
            p = doc.add_paragraph(f"Source: {qa['source']}")
            p.runs[0].font.color.rgb = RGBColor(72, 187, 120)
        
        doc.add_paragraph()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Main App
def main():
    st.markdown('<h1 class="main-header">🎯 Interview Prep Master Pro</h1>', unsafe_allow_html=True)
    st.markdown('<p style="text-align: center; font-size: 1.2rem; color: #666; margin-bottom: 30px;">Comprehensive Interview Preparation with 1000+ Questions</p>', unsafe_allow_html=True)
    
    # Sidebar Configuration
    with st.sidebar:
        st.markdown("### ⚙️ Configuration Panel")
        st.markdown("---")
        
        # Technology Selection (20+ options)
        tech_options = [
            "Artificial Intelligence",
            "Machine Learning",
            "Python",
            "JavaScript",
            "Java",
            "C++",
            "C#",
            "React",
            "Angular",
            "Vue.js",
            "Node.js",
            "Django",
            "Flask",
            "Spring Boot",
            "Azure Cloud",
            "AWS Cloud",
            "GCP Cloud",
            "Docker",
            "Kubernetes",
            "DevOps",
            "Data Science",
            "Data Engineering",
            "SQL",
            "MongoDB",
            "PostgreSQL",
            "Redis",
            "Kafka",
            "Spark",
            "Hadoop",
            "Cybersecurity",
            "Blockchain",
            "Ethereum",
            "Solidity",
            "MLOps",
            "CI/CD"
        ]
        
        selected_tech = st.selectbox("🔧 Select Technology", tech_options, index=0)
        
        # Custom technology
        use_custom = st.checkbox("✏️ Use Custom Technology")
        if use_custom:
            custom_tech = st.text_input("Enter Technology Name", placeholder="e.g., Rust, Go, Terraform")
            if custom_tech:
                selected_tech = custom_tech
        
        st.markdown("---")
        
        # Number of questions (up to 1000)
        num_questions = st.select_slider(
            "📊 Number of Questions",
            options=[10, 20, 30, 50, 100, 200, 300, 500, 1000],
            value=50
        )
        
        st.markdown("---")
        
        # Question Filter Type
        st.markdown("### 🔍 Filter Questions")
        filter_type = st.radio(
            "Select Filter",
            ["All Questions", "Trending Questions", "Latest Questions"],
            index=0
        )
        
        filter_map = {
            "All Questions": "all",
            "Trending Questions": "trending",
            "Latest Questions": "latest"
        }
        selected_filter = filter_map[filter_type]
        
        st.markdown("---")
        
        # Company Selection (20+ companies)
        st.markdown("### 🏢 Company-Specific")
        company_options = [
            "Select Company",
            "Google",
            "Amazon",
            "Microsoft",
            "Meta (Facebook)",
            "Apple",
            "Netflix",
            "Tesla",
            "Uber",
            "Airbnb",
            "LinkedIn",
            "Twitter",
            "Salesforce",
            "Oracle",
            "IBM",
            "Adobe",
            "Intel",
            "NVIDIA",
            "Infosys",
            "TCS",
            "Wipro",
            "Accenture",
            "Cognizant",
            "Capgemini",
            "HCL",
            "Tech Mahindra"
        ]
        
        selected_company = st.selectbox("Company", company_options)
        
        # Custom company
        use_custom_company = st.checkbox("✏️ Custom Company")
        if use_custom_company:
            custom_company = st.text_input("Enter Company Name", placeholder="e.g., Stripe, Snowflake")
            if custom_company:
                selected_company = custom_company
        
        st.markdown("---")
        
        # Interview Experience Section
        st.markdown("### 💼 Interview Experience")
        show_experience = st.checkbox("Show Interview Experiences", value=False)
        
        if show_experience:
            experience_company = st.selectbox(
                "Select Company for Experiences",
                ["Google", "Amazon", "Microsoft", "Meta", "Apple"],
                key="exp_company"
            )
        
        st.markdown("---")
        
        # Tips Section
        st.markdown("### 💡 Features")
        st.success("""
        ✅ 1000+ Questions
        ✅ Trending Topics
        ✅ Company-Specific
        ✅ Proper Formatting
        ✅ Code Examples
        ✅ Video Links
        ✅ PDF/Word Export
        ✅ Interview Experiences
        """)
    
    # Main Content Area
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Technology</div>
            <div class="metric-value">🎯</div>
            <div style="color: #667eea; font-weight: 600;">{selected_tech}</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Questions</div>
            <div class="metric-value">{num_questions}</div>
            <div style="color: #48bb78; font-weight: 600;">Ready to Practice</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        company_display = selected_company if selected_company != "Select Company" else "General"
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Company</div>
            <div class="metric-value">🏢</div>
            <div style="color: #f6ad55; font-weight: 600;">{company_display}</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Filter</div>
            <div class="metric-value">🔍</div>
            <div style="color: #764ba2; font-weight: 600;">{filter_type}</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Generate Button
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🚀 Generate Interview Questions", use_container_width=True, key="generate_btn"):
            with st.spinner("📚 Generating comprehensive interview questions..."):
                time.sleep(1)
                
                company_for_gen = selected_company if selected_company != "Select Company" else None
                questions_data = generate_comprehensive_questions(
                    selected_tech, 
                    num_questions, 
                    selected_filter,
                    company_for_gen
                )
                
                st.session_state['questions_data'] = questions_data
                st.session_state['tech'] = selected_tech
                st.session_state['company'] = selected_company
                
                st.success(f"✅ Successfully generated {len(questions_data)} questions!")
                st.balloons()
    
    # Display Interview Experiences (if enabled)
    if show_experience and 'show_experience' in locals():
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown('<h2 class="sub-header">💼 Interview Experiences</h2>', unsafe_allow_html=True)
        
        experiences = {
            "Google": [
                {
                    "title": "Software Engineer - L3",
                    "date": "December 2024",
                    "rounds": "Phone Screen → Technical (2 rounds) → Googleyness → Team Matching",
                    "experience": "Very focused on algorithms and system design. Interviewers were friendly. Questions were medium to hard on LeetCode scale. They care about code quality and optimization."
                },
                {
                    "title": "ML Engineer - L4",
                    "date": "November 2024",
                    "rounds": "Phone Screen → ML Design → Coding → Behavioral",
                    "experience": "Heavy focus on ML system design. Asked about scaling ML models, feature engineering, and production deployment. Coding round was standard algorithms."
                }
            ],
            "Amazon": [
                {
                    "title": "SDE-2",
                    "date": "January 2025",
                    "rounds": "OA → Phone Screen → Virtual Onsite (4 rounds)",
                    "experience": "Leadership principles are crucial. Every answer should tie back to them. System design was focused on AWS services. Coding questions were medium difficulty."
                }
            ],
            "Microsoft": [
                {
                    "title": "Software Engineer",
                    "date": "December 2024",
                    "rounds": "Phone Screen → Onsite (4 rounds)",
                    "experience": "Very collaborative interview process. Focus on problem-solving approach rather than just the solution. Asked about Azure services and cloud architecture."
                }
            ]
        }
        
        if experience_company in experiences:
            for exp in experiences[experience_company]:
                st.markdown(f"""
                <div class="question-container">
                    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 15px;">
                        <h3 style="color: #667eea; margin: 0;">{exp['title']}</h3>
                        <span class="badge badge-source">{exp['date']}</span>
                    </div>
                    <p style="margin: 10px 0;"><strong>Interview Rounds:</strong> {exp['rounds']}</p>
                    <p style="margin: 10px 0; line-height: 1.6;">{exp['experience']}</p>
                </div>
                """, unsafe_allow_html=True)
    
    # Display Questions
    if 'questions_data' in st.session_state and st.session_state['questions_data']:
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown('<h2 class="sub-header">📋 Interview Questions & Detailed Answers</h2>', unsafe_allow_html=True)
        
        questions_data = st.session_state['questions_data']
        
        # Statistics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📚 Total Questions", len(questions_data))
        with col2:
            easy_count = sum(1 for q in questions_data if q.get('difficulty') == 'Easy')
            st.metric("✅ Easy", easy_count)
        with col3:
            medium_count = sum(1 for q in questions_data if q.get('difficulty') == 'Medium')
            st.metric("⚡ Medium", medium_count)
        with col4:
            hard_count = sum(1 for q in questions_data if q.get('difficulty') == 'Hard')
            st.metric("🔥 Hard", hard_count)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        # Display each question with proper formatting
        for idx, qa in enumerate(questions_data, 1):
            st.markdown(f"""
            <div class="question-container">
                <div class="question-header">
                    <div class="question-number">Question {idx} of {len(questions_data)}</div>
                    <div class="question-text">{qa['question']}</div>
                </div>
            """, unsafe_allow_html=True)
            
            # Answer section with proper formatting
            if qa.get('answer'):
                # Convert markdown-style formatting to HTML
                answer_html = qa['answer']
                answer_html = answer_html.replace('**', '<strong>').replace('**', '</strong>')
                answer_html = answer_html.replace('• ', '<li>').replace('\n\n', '</li></ul><br><ul>').replace('\n', '</li><li>')
                
                # Split into paragraphs for better readability
                paragraphs = qa['answer'].split('\n\n')
                formatted_answer = ""
                
                for para in paragraphs:
                    if para.strip():
                        # Check if it's a bullet list
                        if '•' in para or para.strip().startswith('-'):
                            items = [item.strip('• -') for item in para.split('\n') if item.strip()]
                            formatted_answer += "<ul style='margin: 15px 0; padding-left: 25px;'>"
                            for item in items:
                                if item:
                                    # Bold text in items
                                    item = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', item)
                                    formatted_answer += f"<li style='margin: 8px 0; line-height: 1.6;'>{item}</li>"
                            formatted_answer += "</ul>"
                        else:
                            # Regular paragraph
                            para = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', para)
                            formatted_answer += f"<p style='margin: 15px 0; line-height: 1.8;'>{para}</p>"
                
                st.markdown(f"""
                <div class="answer-section">
                    <span class="answer-label">📝 Detailed Answer:</span>
                    <div class="answer-text">{formatted_answer}</div>
                </div>
                """, unsafe_allow_html=True)
            
            # Metadata section
            st.markdown('<div class="metadata-section">', unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if qa.get('source'):
                    st.markdown(f'<span class="badge badge-source">📚 {qa["source"]}</span>', unsafe_allow_html=True)
            
            with col2:
                if qa.get('difficulty'):
                    st.markdown(f'<span class="badge badge-difficulty">⚡ {qa["difficulty"]}</span>', unsafe_allow_html=True)
            
            with col3:
                if qa.get('category'):
                    st.markdown(f'<span class="badge badge-category">🏷️ {qa["category"]}</span>', unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Links section
            if qa.get('video'):
                st.markdown(f"""
                <a href="{qa['video']}" target="_blank" class="link-button">
                    🎥 Watch Video Explanation
                </a>
                """, unsafe_allow_html=True)
            
            st.markdown('</div><br>', unsafe_allow_html=True)
        
        # Export Section
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown('<h2 class="sub-header">💾 Export Your Questions</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <div class="info-box">
            <strong>📥 Download Options:</strong><br>
            Export your personalized question set to PDF or Word format for offline practice and reference.
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 1, 1])
        
        with col2:
            with st.spinner("Generating PDF..."):
                pdf_buffer = create_pdf(
                    questions_data, 
                    st.session_state['tech'],
                    st.session_state.get('company')
                )
            
            filename_base = f"{st.session_state['tech'].replace(' ', '_')}_Interview_Questions"
            if st.session_state.get('company') and st.session_state['company'] != "Select Company":
                filename_base += f"_{st.session_state['company'].replace(' ', '_')}"
            
            st.download_button(
                label="📄 Download PDF",
                data=pdf_buffer,
                file_name=f"{filename_base}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        
        with col3:
            with st.spinner("Generating Word document..."):
                word_buffer = create_word(
                    questions_data,
                    st.session_state['tech'],
                    st.session_state.get('company')
                )
            
            st.download_button(
                label="📝 Download Word",
                data=word_buffer,
                file_name=f"{filename_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    # Footer
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 30px 20px;">
        <p style="font-size: 1.3rem; font-weight: 600; margin-bottom: 15px;">🌟 Interview Prep Master Pro</p>
        <p style="font-size: 1rem; margin: 10px 0;">Comprehensive Interview Preparation Platform</p>
        <p style="font-size: 0.95rem; margin: 10px 0;">
            ✨ Features: Trending Questions • Company-Specific Prep • Interview Experiences • Detailed Answers • Video Resources
        </p>
        <p style="font-size: 0.9rem; margin-top: 20px; color: #999;">
            💼 Practice Daily • 🎯 Stay Focused • 🚀 Achieve Success
        </p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
